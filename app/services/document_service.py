from typing import List, Dict, Any, Optional
import io
from pathlib import Path
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from bs4 import BeautifulSoup
import markdown
from docx import Document as DocxDocument
from app.services.ocr_service import ocr_service


class DocumentService:
    """Service for processing and chunking documents for vector storage"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.html', '.htm', '.docx'}

    @staticmethod
    def extract_text_from_pdf(file_content: bytes, force_ocr: bool = False) -> str:
        """
        Extract text from PDF file, using OCR if necessary

        Args:
            file_content: Binary content of the PDF file
            force_ocr: If True, always use OCR regardless of text detection

        Returns:
            Extracted text content
        """
        try:
            # Use OCR service which automatically detects if OCR is needed
            return ocr_service.extract_text_from_pdf_with_ocr(file_content, force_ocr)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_html(file_content: bytes) -> str:
        """Extract text from HTML file"""
        try:
            html_text = file_content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_text, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text
        except Exception as e:
            raise ValueError(f"Failed to extract text from HTML: {str(e)}")

    @staticmethod
    def extract_text_from_markdown(file_content: bytes) -> str:
        """Extract text from Markdown file"""
        try:
            md_text = file_content.decode('utf-8', errors='ignore')
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        except Exception as e:
            raise ValueError(f"Failed to extract text from Markdown: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            # Load document from bytes
            doc = DocxDocument(io.BytesIO(file_content))

            # Extract text from all paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]

            # Also extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    table_texts.append(row_text)

            # Combine paragraphs and tables
            full_text = '\n'.join(paragraphs)
            if table_texts:
                full_text += '\n\nTables:\n' + '\n'.join(table_texts)

            return full_text
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text(filename: str, file_content: bytes) -> str:
        """
        Extract text from a file based on its extension

        Args:
            filename: Name of the file
            file_content: Binary content of the file

        Returns:
            Extracted text content
        """
        extension = Path(filename).suffix.lower()

        if extension == '.pdf':
            return DocumentService.extract_text_from_pdf(file_content)
        elif extension in ['.html', '.htm']:
            return DocumentService.extract_text_from_html(file_content)
        elif extension == '.md':
            return DocumentService.extract_text_from_markdown(file_content)
        elif extension == '.txt':
            return DocumentService.extract_text_from_txt(file_content)
        elif extension == '.docx':
            return DocumentService.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Split text into chunks using RecursiveCharacterTextSplitter

        Args:
            text: Text to split
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks
            metadata: Optional metadata to attach to each chunk

        Returns:
            List of Document objects with chunked text and metadata
        """
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        chunks = text_splitter.split_text(text)

        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy() if metadata else {}
            doc_metadata['chunk_index'] = i
            doc_metadata['chunk_count'] = len(chunks)

            documents.append(Document(
                page_content=chunk,
                metadata=doc_metadata
            ))

        return documents

    @staticmethod
    def process_document(
        filename: str,
        file_content: bytes,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[Document]:
        """
        Process a document: extract text and split into chunks

        Args:
            filename: Name of the file
            file_content: Binary content of the file
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks

        Returns:
            List of Document objects with chunked text and metadata
        """
        # Validate file type
        extension = Path(filename).suffix.lower()
        if extension not in DocumentService.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(DocumentService.SUPPORTED_EXTENSIONS)}"
            )

        # Extract text
        text = DocumentService.extract_text(filename, file_content)

        if not text.strip():
            raise ValueError("No text content found in document")

        # Prepare metadata
        metadata = {
            'source': filename,
            'file_type': extension[1:],  # Remove the dot
            'uploaded_at': datetime.utcnow().isoformat(),
        }

        # Chunk text
        documents = DocumentService.chunk_text(
            text=text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            metadata=metadata
        )

        return documents


# Singleton instance
document_service = DocumentService()
